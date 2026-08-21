#!/usr/bin/env python3
"""Generate synthetic test case evidence files (TXT + PDF + DOCX) for 3 domains."""

import os
from pathlib import Path
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document

BASE = Path(__file__).parent / "sample_cases"


def make_pdf(path, title, paragraphs):
    """Create a simple PDF with title and paragraphs."""
    doc = SimpleDocTemplate(str(path), pagesize=A4, margins=2 * cm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="BodyDE", parent=styles["Normal"], fontSize=10, leading=14))
    story = [Paragraph(title, styles["Title"]), Spacer(1, 0.5 * cm)]
    for p in paragraphs:
        story.append(Paragraph(p, styles["BodyDE"]))
        story.append(Spacer(1, 0.2 * cm))
    doc.build(story)


def make_docx(path, title, paragraphs):
    """Create a simple DOCX with title and paragraphs."""
    doc = Document()
    doc.add_heading(title, level=0)
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))


# ═══════════════════════════════════════════════════════════════
# CASE 1: LEGAL — Arbeitsplatzunfall, Zivilprozess
# ═══════════════════════════════════════════════════════════════

LEGAL_TXT = """\
Klagschrift — Arbeitsplatzunfall

Gerichtsaktenzeichen: Az. 2 O 415/2026
Gericht: Landgericht Berlin — Zivilkammer 4
Datum: 14.03.2026

Kläger: Thomas Becker
Geburtsdatum: 22.07.1981
Anschrift: Goethestr. 12, 10625 Berlin
Steuernummer: BCKRTH81L22G123X

Beklagte: Bauunternehmen Himmel & Erde GmbH
Geschäftsführer: Wolfgang Schmidt
Anschrift: Industrieallee 78, 12359 Berlin
Steuernummer: HEGMBH78B1234Y

== SACHVERHALT ==

Am 10.01.2026 um 09:15 Uhr ereignete sich auf der Baustelle
Linden-Center, Frankfurt/Oder ein Arbeitsunfall.
Der Kläger war als Dachdecker auf dem Gerüst in ca. 8 Metern Höhe tätig.
Das Gerüst wies an der Ostseite keine ausreichenden Seitenschutzvorrichtungen auf.
Der Kläger verlor das Gleichgewicht und stürzte ca. 6 Meter in die Tiefe.

== SCHADENSPOSITIONEN ==

1. Krankenhausbehandlung: EUR 28.450,00
   - Notoperation, 5 Tage Intensivstation, 12 Tage Normalstation
   - Unfallkrankenhaus Berlin-Marzahn, Datum: 10.01.2026 bis 27.01.2026

2. Verdienstausfall: EUR 15.300,00
   - Nettoeinkommen vor Unfall: EUR 3.500,00 monatlich
   - Lohnfortzahlung: 6 Wochen (bis 21.02.2026)
   - Ausfall danach: 10.02.2026 bis 31.05.2026

3. Schmerzensgeld: EUR 50.000,00
   - geschätzter Anspruch nach Dassler-Beck-Eschertabelle
   - Schwere Verletzung: Beckenbruch, Rippenbrüche, Schädelhirntrauma

4. Haushaltsfuehrungsschaden: EUR 4.200,00
   - 14 Wochen eingeschränkte Haushaltsfuehrung
   - Haushaltshilfe: EUR 300,00 pro Woche

Gesamtschaden: EUR 97.950,00

== RECHTLICHE WÜRDIGUNG ==

Der Kläger macht Schadensersatzansprüche nach § 823 Abs. 1 BGB,
§§ 618, 249 ff. BGB geltend.
Die Beklagte hat ihre Verkehrspflichten aus § 618 BGB verletzt,
indem das Gerüst nicht mit ausreichendem Seitenschutz ausgestattet war.
Die Kausalität zwischen Pflichtverletzung und Schaden ist gegeben.
Die Beweislast für die Einhaltung der Sicherheitsvorschriften
trifft die Beklagte.

Zeugen:
  Zeuge 1: Petra Hoffmann, Bauarbeiterin, Datum der Aussage: 15.02.2026
  Zeuge 2: Frank Müller, Gerüstbauer, Datum der Aussage: 18.02.2026

Beweismittel:
  - Gerichtliches Gutachten (siehe separate Datei)
  - Zeugenaussagen (siehe separate Datei)
  - Krankenhausbericht (Unfallkrankenhaus Berlin-Marzahn)
  - Lohnabrechnungen Oktober 2025 bis Dezember 2025
  - Fotos der Baustelle (Blatt 1-12 der Akte)
"""

LEGAL_GUTACHTEN_PARAS = [
    "Gerichtliches Sachverstaendigengutachten — Az. 2 O 415/2026",
    "Gericht: Landgericht Berlin — Zivilkammer 4",
    "Gutachter: Dr.-Ing. Klaus Rentrop, Diplom-Ingenieur fuer Arbeitssicherheit",
    "Gutachtennummer: GUT-2026-0042",
    "Datum: 20.03.2026",
    "",
    "Auftrag: Pruefung der Sicherheitsvorkehrungen auf der Baustelle Linden-Center, Frankfurt/Oder, zum Zeitpunkt des Unfalls am 10.01.2026 um 09:15 Uhr.",
    "",
    "1. Feststellungen vor Ort",
    "Das Gerüst an der Ostseite des Gebaeudes wies zum Unfallzeitpunkt keinen Seitenschutz auf. Die Schutzplanken waren auf einer Laenge von ca. 4 Metern entfernt. Der Geruestbelag war mit Schnee und Eis bedeckt. Die Rutschgefahr war erheblich. Ein Sicherungsgurt wurde dem Klaeger nicht zur Verfuegung gestellt.",
    "",
    "2. Bewertung",
    "Nach der Unfallverhuetungsvorschrift BGV C22 (Gerueste) muss ein Geruest in mehr als 3 Metern Hoehe einen vollstaendigen Seitenschutz aufweisen. Die Beklagte hat diese Vorschrift deutlich verletzt. Der Seitenschutz war nicht nur unvollstaendig, sondern auf einer Laenge von 4 Metern komplett entfernt. Dies stellt eine schwerwiegende Verletzung der Verkehrspflicht dar.",
    "",
    "3. Kausalitaet",
    "Der Sturz des Klaegers ist mit ueberwiegender Wahrscheinlichkeit auf das Fehlen des Seitenschutzes und die Eisbildung zurueckzufuehren. Ein vollstaendiger Seitenschutz haette den Sturz verhindert oder zumindest deutlich abgemildert. Die Kausalitaet wird mit einer Wahrscheinlichkeit von ueber 90% bejaht.",
    "",
    "4. Schadensschaetzung",
    "Die Heilbehandlungskosten in Hoehe von EUR 28.450,00 sind durch die Rechnungen des Unfallkrankenhauses Berlin-Marzahn belegt. Der Verdienstausfall von EUR 15.300,00 ist plausibel berechnet. Das Schmerzensgeld von EUR 50.000,00 erscheint angemessen, jedoch eher am unteren Rand. Ein Betrag von EUR 60.000,00 bis EUR 75.000,00 waere ebenfalls vertretbar.",
    "",
    "5. Zusammenfassung",
    "Die Beklagte hat ihre Verkehrspflichten aus § 618 BGB und der BGV C22 verletzt. Die Kausalitaet zwischen Pflichtverletzung und Unfall ist mit hoher Wahrscheinlichkeit gegeben. Die Schadenspositionen sind im Wesentlichen belegt, das Schmerzensgeld duerfte jedoch hoeher ausfallen.",
]

LEGAL_ZEUGEN_PARAS = [
    "Zeugenaussage — Az. 2 O 415/2026",
    "Gericht: Landgericht Berlin — Zivilkammer 4",
    "Datum: 15.02.2026",
    "",
    "Zeugin: Petra Hoffmann",
    "Beruf: Bauarbeiterin (Maurerin)",
    "Anschrift: Karl-Marx-Allee 55, 10243 Berlin",
    "",
    "Eidesstattliche Versicherung:",
    "Ich, Petra Hoffmann, erklaere hiermit an Eides statt:",
    "",
    "Ich war am 10.01.2026 ab 07:30 Uhr auf der Baustelle Linden-Center in Frankfurt/Oder im Einsatz. Gegen 09:00 Uhr befand ich mich im Erdgeschoss und hoerte einen lauten Ruf. Als ich nach draussen lief, sah ich Herrn Becker am Boden liegen. Er war bewusstlos.",
    "",
    "Ich habe das Geruest am Vortag, den 09.01.2026, bereits bemerkt, dass der Seitenschutz an der Ostseite fehlte. Ich habe Herrn Schmidt, den Bauleiter, darauf hingewiesen. Er sagte, die Planken wuerden am Nachmittag montiert. Dies ist jedoch nicht geschehen.",
    "",
    "Die Bedingungen am Unfalltag waren frostig. Das Geruest war vereist. Ich selbst hatte grosse Angst, auf das Geruest zu steigen, und habe im Erdgeschoss gearbeitet.",
    "",
    "Ich bestaetige, dass Herr Becker zum Unfallzeitpunkt keinen Sicherungsgurt trug. Dieser wurde uns auch nicht zur Verfuegung gestellt. Es war auch kein Sicherheitsbeauftragter anwesend.",
    "",
    "Vor mir ist dies bereits am 05.12.2025 passiert, dass ein Geruestbauer fast abgestuerzt waere. Damals wurde auch nichts unternommen.",
    "",
    "Ort, Datum: Berlin, 15.02.2026",
    "Unterschrift: Petra Hoffmann",
]


# ═══════════════════════════════════════════════════════════════
# CASE 2: CYBER — Ransomware-Angriff
# ═══════════════════════════════════════════════════════════════

CYBER_TXT = """\
Incident Response Report — Ransomware-Angriff

Vorfall-ID: INC-2026-0312
Unternehmen: TechNova Solutions GmbH
Anschrift: Silicon Valley 5, 41769 Viersen
IT-Leiter: Stefan Weber
Telefon: +49 2162 1234567
Email: s.weber@technova.de

== ZUSAMMENFASSUNG ==

Am 18.03.2026 um 03:47:32 UTC wurde ein Ransomware-Angriff auf die
Infrastruktur der TechNova Solutions GmbH erkannt.
Der Angriff begann am 17.03.2026 um 22:14:08 UTC mit einem
Phishing-Email an den Mitarbeiter Andreas Klein.
Die Ransomware verschluesselte 4 Server und 12 Workstations.
Der Schaden wird auf EUR 250.000,00 geschaetzt
(Stillstandskosten: EUR 180.000,00, Wiederherstellung: EUR 70.000,00).

== BETROFFENE SYSTEME ==

Server:
  SRV-DC01    — Domain Controller (Windows Server 2019)
  SRV-FS02    — Fileserver (Windows Server 2022)
  SRV-EX03    — Exchange Server (Windows Server 2019)
  SRV-APP04   — Application Server (Ubuntu 22.04)

Workstations:
  WS-HR-01 bis WS-HR-04    — HR-Abteilung
  WS-FIN-01 bis WS-FIN-04  — Finanzabteilung
  WS-DEV-01 bis WS-DEV-04  — Entwicklung

== ANGRIFFSZEITLEISTE ==

2026-03-17T22:14:08Z  Phishing-Email empfangen (Andreas Klein)
2026-03-17T22:31:55Z  Oeffnen des Malware-Anhangs auf WS-HR-01
2026-03-17T22:35:12Z  C2-Verbindung zu 185.220.101.34
2026-03-17T23:02:44Z  Lateral Movement: WS-HR-01 -> SRV-DC01
2026-03-18T01:15:30Z  Credential Dumping auf SRV-DC01
2026-03-18T02:44:18Z  Deployment der Ransomware auf SRV-FS02
2026-03-18T03:12:07Z  Verschluesselung beginnt auf SRV-FS02
2026-03-18T03:47:32Z  Erkennung durch Monitoring-Alert
2026-03-18T04:01:55Z  Netzwerkisolation aller Systeme

== INDIKATOREN DER KOMPROMITTIERUNG (IOCs) ==

IP-Adressen:
  185.220.101.34    — C2-Server (Tor Exit Node)
  45.133.1.197      — Phishing-Domain IP
  91.218.114.52     — Exfiltration-Server

Datei-Hashes (SHA-256):
  3a7b...e9f2       — ransomware.exe
  8c2d...4a1b       — cobalt_strike_beacon.dll

Mutationsdatum: 2026-03-18T03:12:07Z

== ERKENNTE SCHWACHSTELLEN ==

1. Keine MFA auf RDP-Verbindungen (Port 3389 offen im Internet)
2. Veraltete Exchange-Server (CVE-2026-1234 nicht gepatcht)
3. Fehlende Netzwerksegmentierung zwischen HR und Servernetz
4. Keine Application Whitelisting auf Workstations
5. Virenschutz-Definitionen 14 Tage veraltet

== MASSNAHMEN ==

1. Netzwerkisolation aller betroffenen Systeme
2. Abschaltung von SRV-DC01, SRV-FS02, SRV-EX03, SRV-APP04
3. Forensische Images aller betroffenen Systeme
4. Reset aller Domain-Accounts und Service-Passwoerter
5. Wiederherstellung aus Offline-Backups (Stand: 16.03.2026)
6. Meldung an BSI und Landesdatenschutzbeauftragten
"""

CYBER_LOG_PARAS = [
    "Server-Log — SRV-FS02 (Fileserver)",
    "System: SRV-FS02, IP: 192.168.1.50",
    "Betriebssystem: Windows Server 2022",
    "",
    "2026-03-18T02:44:18Z — Event ID 4688: Neuer Prozess erstellt",
    "  Account: TECHNOVA\\\\admin_klein",
    "  Prozess: C:\\\\Users\\\\admin_klein\\\\AppData\\\\Local\\\\Temp\\\\ransomware.exe",
    "  Befehlszeile: ransomware.exe --encrypt-all --extensions .docx,.xlsx,.pdf,.txt",
    "",
    "2026-03-18T02:48:33Z — Event ID 4688: vssadmin.exe",
    "  Account: TECHNOVA\\\\admin_klein",
    "  Befehlszeile: vssadmin delete shadows /all /quiet",
    "",
    "2026-03-18T03:12:07Z — Event ID 4663: Massenhafte Dateizugriffe",
    "  Pfad: \\\\\\\\SRV-FS02\\\\Freigabe\\\\Personal\\\\",
    "  Zugriffstyp: DELETE + WRITE",
    "  Anzahl betroffener Dateien: 14.832",
    "",
    "2026-03-18T03:14:22Z — Firewall-Alert: Ausgehende Verbindung",
    "  Quelle: 192.168.1.50:49152",
    "  Ziel: 91.218.114.52:443",
    "  Protokoll: HTTPS",
    "  Bytes gesendet: 2.3 GB",
    "  Dauer: 33 Minuten",
    "",
    "2026-03-18T03:47:32Z — Monitoring-Alert ausgelöst",
    "  Alert: Mass file encryption detected on SRV-FS02",
    "  Schweregrad: CRITICAL",
    "  Trigger: Datei-Integritätsmonitoring",
    "",
    "2026-03-18T03:52:15Z — Firewall-Log: Eingehende Verbindung blockiert",
    "  Quelle: 185.220.101.34",
    "  Ziel: 192.168.1.50:3389",
    "  Protokoll: RDP",
    "  Aktion: BLOCK (nach Isolation)",
]

CYBER_FW_PARAS = [
    "Firewall-Log — PF-Firewall (FW-EDGE-01)",
    "System: FW-EDGE-01, IP: 192.168.1.1",
    "Datum: 17.03.2026 bis 18.03.2026",
    "",
    "2026-03-17T22:14:08Z — ALLOW SMTP",
    "  Quelle: 45.133.1.197:587",
    "  Ziel: 192.168.1.10:25",
    "  Protokoll: SMTP",
    "  Notiz: Email an a.klein@technova.de empfangen",
    "",
    "2026-03-17T22:35:12Z — ALLOW HTTPS",
    "  Quelle: 192.168.1.31:49152",
    "  Ziel: 185.220.101.34:443",
    "  Protokoll: HTTPS",
    "  Notiz: C2-Verbindung von WS-HR-01",
    "  WS-HR-01 IP: 192.168.1.31",
    "",
    "2026-03-17T23:02:44Z — ALLOW RDP",
    "  Quelle: 192.168.1.31:49153",
    "  Ziel: 192.168.1.10:3389",
    "  Protokoll: RDP",
    "  Notiz: Lateral Movement WS-HR-01 -> SRV-DC01",
    "  SRV-DC01 IP: 192.168.1.10",
    "",
    "2026-03-18T03:14:22Z — ALLOW HTTPS (suspicious)",
    "  Quelle: 192.168.1.50:49152",
    "  Ziel: 91.218.114.52:443",
    "  Protokoll: HTTPS",
    "  Bytes: 2.3 GB",
    "  Notiz: Datenexfiltration von SRV-FS02",
    "",
    "2026-03-18T04:01:55Z — BLOCK ALL",
    "  Quelle: 0.0.0.0",
    "  Ziel: 192.168.1.0/24",
    "  Protokoll: ALL",
    "  Notiz: Notfall-Isolation durch IT-Leiter Stefan Weber",
    "",
    "2026-03-18T04:15:30Z — BLOCK RDP",
    "  Quelle: 185.220.101.34",
    "  Ziel: 192.168.1.50:3389",
    "  Protokoll: RDP",
    "  Notiz: Versuchter erneuter Zugriff nach Isolation",
]


# ═══════════════════════════════════════════════════════════════
# CASE 3: FINANCIAL — Betrugsverdacht Buchhaltung
# ═══════════════════════════════════════════════════════════════

FINANCIAL_TXT = """\
Pruefbericht — Betrugsverdacht Buchhaltung

Aktenzeichen: PRF-2026-0078
Pruefdatum: 12.06.2026
Pruefer: Sabine Vogt, Wirtschaftsprueferin
Mandant: Müller Elektrotechnik GmbH & Co. KG
Anschrift: Gewerbestraße 33, 85055 Ingolstadt
Steuernummer: MUELEK81B2345K

Beschuldigter: Jürgen Wagner
Position: Buchhalter (seit 01.03.2019)
Geburtsdatum: 03.11.1975
Anschrift: Schillerweg 7, 85072 Eichstätt
Steuernummer: WGNJGN75R03E084Z

== VERDACHTSMOMENTE ==

1. Auffällige Ueberweisungen an unbekannte IBANs
2. Buchungen ohne corresponding Beleg
3. Rechnungen mit identischen Betraegen von unterschiedlichen Rechnungsstellern
4. Buchungen ausserhalb der Geschäftszeiten
5. Haeufige Korrekturbuchungen durch Herrn Wagner

== AUFFAELLIGE BUCHUNGEN ==

Buchung 1:
  Datum: 15.01.2026
  Betrag: EUR 12.500,00
  Von: DE89370400440532013000 (Müller Elektrotechnik Geschäftskonto)
  An: DE71500105179423614829 (unbekannt, TechConsult Pro GmbH)
  Buchungstext: IT-Beratungsleistung Q4/2025
  Beleg: Rechnung RE-2026-0042 (vorhanden, aber unplausicbel)
  Bemerkung: TechConsult Pro GmbH nicht im Handelsregister auffindbar

Buchung 2:
  Datum: 22.01.2026
  Betrag: EUR 12.500,00
  Von: DE89370400440532013000 (Müller Elektrotechnik Geschäftskonto)
  An: DE48500105169874125369 (unbekannt, Digital Solutions AG)
  Buchungstext: Wartungsvertrag IT-Infrastruktur
  Beleg: Rechnung RE-2026-0051 (vorhanden, identischer Betrag wie Buchung 1)

Buchung 3:
  Datum: 05.02.2026
  Betrag: EUR 8.750,00
  Von: DE89370400440532013000 (Müller Elektrotechnik Geschäftskonto)
  An: DE71500105179423614829 (gleiche IBAN wie Buchung 1!)
  Buchungstext: Software-Lizenzierung
  Beleg: keine Rechnung vorhanden

Buchung 4:
  Datum: 18.02.2026
  Betrag: EUR 12.500,00
  Von: DE89370400440532013000 (Müller Elektrotechnik Geschäftskonto)
  An: DE38200700200569874123 (unbekannt, Cloud Services UG)
  Buchungstext: Cloud-Migration Januar 2026
  Beleg: Rechnung RE-2026-0078 (vorhanden, wieder EUR 12.500,00)

Buchung 5:
  Datum: 03.03.2026
  Betrag: EUR 5.300,00
  Von: DE89370400440532013000 (Müller Elektrotechnik Geschäftskonto)
  An: DE71500105179423614829 (gleiche IBAN wie Buchung 1 und 3!)
  Buchungstext: IT-Schulung Mitarbeitende
  Beleg: Rechnung RE-2026-0095

Gesamtauffälliges Volumen: EUR 51.550,00

== WEITERE AUFFAELLIGKEITEN ==

- Herr Wagner hat alle 5 Buchungen persönlich freigegeben
- Die Freigaben erfolgten teilweise ausserhalb der Geschäftszeiten (20:47, 21:15, 06:03)
- Die IBAN DE71500105179423614829 erscheint 3x — die zugehörige
  Firma TechConsult Pro GmbH ist nicht im Handelsregister eingetragen
- Die Rechnungen RE-2026-0042 und RE-2026-0051 haben identische Beträge
  und aehnlichen Wortlaut, stammen aber von unterschiedlichen angeblichen Firmen
- Die Umsatzsteuer-IDs auf den Rechnungen sind nicht validierbar

== EMPFEHLUNG ==

1. Sofortige Suspendierung von Herrn Wagner
2. Forensische Auswertung aller Buchungen der letzten 24 Monate
3. Pruefung aller Rechnungen von TechConsult Pro GmbH und Digital Solutions AG
4. Strafanzeige wegen Betrugsverdacht § 263 StGB
5. Meldung an die Steuerfahndung
"""

FINANCIAL_BANK_PARAS = [
    "Bankauszug — Müller Elektrotechnik GmbH & Co. KG",
    "Bank: Sparkasse Ingolstadt",
    "IBAN: DE89370400440532013000",
    "Kontonummer: 0532013000",
    "Bankleitzahl: 72140043",
    "Zeitraum: Januar 2026 bis Maerz 2026",
    "",
    "Saldo per 01.01.2026: EUR 487.234,56",
    "",
    "Buchungsdatum 15.01.2026:",
    "  Ueberweisung: EUR 12.500,00",
    "  Empfaenger: TechConsult Pro GmbH",
    "  IBAN: DE71500105179423614829",
    "  BIC: PBNKDEFF",
    "  Verwendungszweck: IT-Beratungsleistung Q4/2025 Rechnung RE-2026-0042",
    "  Saldo nach Buchung: EUR 474.734,56",
    "",
    "Buchungsdatum 22.01.2026:",
    "  Ueberweisung: EUR 12.500,00",
    "  Empfaenger: Digital Solutions AG",
    "  IBAN: DE48500105169874125369",
    "  BIC: WELADEXX",
    "  Verwendungszweck: Wartungsvertrag IT-Infrastruktur Rechnung RE-2026-0051",
    "  Saldo nach Buchung: EUR 462.234,56",
    "",
    "Buchungsdatum 05.02.2026:",
    "  Ueberweisung: EUR 8.750,00",
    "  Empfaenger: TechConsult Pro GmbH",
    "  IBAN: DE71500105179423614829",
    "  BIC: PBNKDEFF",
    "  Verwendungszweck: Software-Lizenzierung",
    "  Saldo nach Buchung: EUR 453.484,56",
    "",
    "Buchungsdatum 18.02.2026:",
    "  Ueberweisung: EUR 12.500,00",
    "  Empfaenger: Cloud Services UG",
    "  IBAN: DE38200700200569874123",
    "  BIC: DEUTDEFF",
    "  Verwendungszweck: Cloud-Migration Januar 2026 Rechnung RE-2026-0078",
    "  Saldo nach Buchung: EUR 440.984,56",
    "",
    "Buchungsdatum 03.03.2026:",
    "  Ueberweisung: EUR 5.300,00",
    "  Empfaenger: TechConsult Pro GmbH",
    "  IBAN: DE71500105179423614829",
    "  BIC: PBNKDEFF",
    "  Verwendungszweck: IT-Schulung Mitarbeitende Rechnung RE-2026-0095",
    "  Saldo nach Buchung: EUR 435.684,56",
    "",
    "Saldo per 31.03.2026: EUR 435.684,56",
    "Summe auffaelliger Ueberweisungen: EUR 51.550,00",
]

FINANCIAL_RECHNUNG_PARAS = [
    "Rechnung RE-2026-0042",
    "",
    "Rechnungssteller: TechConsult Pro GmbH",
    "Anschrift: Berliner Allee 15, 40212 Duesseldorf",
    "USt-IdNr.: DE321456789",
    "Steuernummer: 133/5821/6421",
    "",
    "Rechnungsempfaenger: Mueller Elektrotechnik GmbH & Co. KG",
    "Anschrift: Gewerbestrasse 33, 85055 Ingolstadt",
    "",
    "Rechnungsdatum: 10.01.2026",
    "Leistungsdatum: 01.12.2025 bis 31.12.2025",
    "Rechnungsnummer: RE-2026-0042",
    "",
    "Leistungsbeschreibung:",
    "  IT-Beratungsleistung Q4/2025",
    "  - Strategieberatung IT-Architektur: EUR 5.000,00",
    "  - Systemmigration Konzept: EUR 4.500,00",
    "  - Projektmanagement: EUR 3.000,00",
    "  Zwischensumme netto: EUR 12.500,00",
    "  Umsatzsteuer 19%: EUR 2.375,00",
    "  Rechnungsgesamtbetrag: EUR 14.875,00",
    "",
    "Zahlungsbedingungen: 14 Tage netto",
    "Bankverbindung: IBAN DE71500105179423614829, BIC PBNKDEFF",
    "",
    "Hinweis: Die Rechnung wurde von Jürgen Wagner am 15.01.2026 freigegeben.",
    "  Freigabezeit: 20:47 Uhr (ausserhalb der Geschäftszeiten)",
    "",
    "AUFFAELLIGKEITEN:",
    "  - TechConsult Pro GmbH ist nicht im Handelsregister Duesseldorf eingetragen",
    "  - Die USt-IdNr. DE321456789 ist nicht validierbar",
    "  - Die Leistungsbeschreibung ist unkonkret (keine Stunden, keine Projektphasen)",
    "  - Die Bankverbindung fuehrt zu einer Postbank-Filiale in Muenchen",
    "    (nicht in Duesseldorf, wo die Firma angeblich sitzt)",
    "  - Der Rechnungsbetrag von EUR 14.875,00 wurde ueberwiesen,",
    "    der Buchungstext enthaelt jedoch nur EUR 12.500,00",
    "    (Differenz von EUR 2.375,00 = Umsatzsteuer)",
]


# ═══════════════════════════════════════════════════════════════
# GENERATE ALL FILES
# ═══════════════════════════════════════════════════════════════

def main():
    # --- Legal ---
    legal = BASE / "legal_001" / "evidence"
    (legal / "klageschrift.txt").write_text(LEGAL_TXT, encoding="utf-8")
    make_pdf(legal / "gerichtsgutachten.pdf", "Gerichtliches Sachverstaendigengutachten", LEGAL_GUTACHTEN_PARAS)
    make_docx(legal / "zeugenaussage.docx", "Zeugenaussage — Petra Hoffmann", LEGAL_ZEUGEN_PARAS)

    # --- Cyber ---
    cyber = BASE / "cyber_001" / "evidence"
    (cyber / "incident_report.txt").write_text(CYBER_TXT, encoding="utf-8")
    make_pdf(cyber / "server_log.pdf", "Server-Log — SRV-FS02", CYBER_LOG_PARAS)
    make_docx(cyber / "firewall_log.docx", "Firewall-Log — FW-EDGE-01", CYBER_FW_PARAS)

    # --- Financial ---
    financial = BASE / "financial_001" / "evidence"
    (financial / "pruefbericht.txt").write_text(FINANCIAL_TXT, encoding="utf-8")
    make_pdf(financial / "bankauszug.pdf", "Bankauszug — Müller Elektrotechnik", FINANCIAL_BANK_PARAS)
    make_docx(financial / "rechnung.docx", "Rechnung RE-2026-0042", FINANCIAL_RECHNUNG_PARAS)

    print("All evidence files generated successfully.")
    for case in ["legal_001", "cyber_001", "financial_001"]:
        d = BASE / case / "evidence"
        files = sorted(f.name for f in d.iterdir())
        print(f"  {case}: {files}")


if __name__ == "__main__":
    main()